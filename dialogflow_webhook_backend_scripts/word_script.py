#this script basically has funtions that will manipulate ms word
#module named docx was used which has functions that help manipulate word and docx files
import docx
from docx import Document

#this is used to import function open_document
#this function opens the file passed in arguments using system module of python
from system_commands import open_document

path = 'Desktop/Internet of Things/project/files/test_data.docx'

def create_document():
    #following command creates a new document at the specifies path
    document = Document()
    document.add_heading("new document")
    document.save(path)

def write_into_document(data__recieved):
    data_to_type = (data__recieved['result']['parameters']['any'])
    print(data_to_type)
    document = docx.Document(path)
    document.add_paragraph(data_to_type)
    document.save(path)
    open_document(path)

def append_into_document(data_recieved):
    data_to_type = data_recieved['result']['parameters']['any']
    print(data_to_type)
    document = docx.Document(path)
    document.add_paragraph(data_to_type)
    document.save(path)
    open_document(path)

def add_heading_into_document(data_recieved):
    data_to_type = data_recieved['result']['parameters']['any']
    print(data_to_type)
    document = docx.Document(path)
    document.add_heading(data_to_type)
    document.save(path)
    open_document(path)

def first_paragraph_bold():
    document = docx.Document(path)
    first_paragraph = document.paragraphs[0]
    runs = first_paragraph.runs
    runs[0].bold = True
    document.save(path)
    open_document(path)

def first_paragraph_italic():
    document = docx.Document(path)
    first_paragraph = document.paragraphs[0]
    runs = first_paragraph.runs
    runs[0].italic = True
    document.save(path)
    open_document(path)

def first_paragraph_underline():
    document = docx.Document(path)
    first_paragraph = document.paragraphs[0]
    runs = first_paragraph.runs
    runs[0].underline = True
    document.save(path)
    open_document(path)

